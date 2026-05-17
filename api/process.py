import base64
import io
import logging
import os
import re

from docx import Document
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from io import BytesIO, StringIO
from pipeline import process_data
import pandas as pd
from PyPDF2 import PdfReader

app = Flask(__name__)
CORS(app)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {"csv", "xlsx", "docx", "pdf"}
PUBLIC_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public"))


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_csv(file_stream):
    file_stream.seek(0)
    content = file_stream.read().decode("utf-8", errors="replace")
    return pd.read_csv(StringIO(content))


def parse_excel(file_stream):
    file_stream.seek(0)
    return pd.read_excel(file_stream, engine="openpyxl")


def parse_docx(file_stream):
    file_stream.seek(0)
    doc = Document(file_stream)
    if not doc.tables:
        raise ValueError("No table found in Word document.")

    table = doc.tables[0]
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])

    if len(rows) < 2:
        raise ValueError("Word file must contain a header row and at least one student row.")

    return pd.DataFrame(rows[1:], columns=rows[0])


def parse_pdf(file_stream):
    file_stream.seek(0)
    reader = PdfReader(file_stream)
    raw_text = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        raw_text.append(page_text)

    raw = "\n".join(raw_text).strip()
    if not raw:
        raise ValueError("PDF file contains no readable text.")

    lines = [line.strip() for line in raw.splitlines() if line.strip()]
    if len(lines) < 2:
        raise ValueError("PDF file must contain headers and at least one row.")

    header = re.split(r"\s{2,}|\t|,", lines[0])
    rows = []
    for line in lines[1:]:
        cells = re.split(r"\s{2,}|\t|,", line)
        if len(cells) == len(header):
            rows.append(cells)
            continue
        parts = re.split(r"\s+", line)
        if len(parts) == len(header):
            rows.append(parts)
            continue
        if len(parts) > len(header) and len(header) >= 2:
            diff = len(parts) - len(header) + 1
            merged_name = " ".join(parts[:diff])
            rows.append([merged_name] + parts[diff:])

    if not rows:
        raise ValueError("Unable to parse table data from PDF. Ensure the PDF contains a simple student table.")

    return pd.DataFrame(rows, columns=header)


def load_dataframe(file_storage):
    filename = file_storage.filename
    extension = filename.rsplit(".", 1)[1].lower()
    stream = file_storage.stream

    if extension == "csv":
        df = parse_csv(stream)
    elif extension == "xlsx":
        df = parse_excel(stream)
    elif extension == "docx":
        df = parse_docx(stream)
    elif extension == "pdf":
        df = parse_pdf(stream)
    else:
        raise ValueError("Unsupported file type.")

    return df


def normalize_dataframe(df):
    df = df.copy()
    df.columns = [str(col).strip() for col in df.columns]
    if "Name" in df.columns:
        df["Name"] = df["Name"].astype(str).str.strip()

    extra_columns = [col for col in df.columns if col not in ["Name", "Total", "Average", "Grade", "Rank"]]
    for col in extra_columns:
        df[col] = pd.to_numeric(df[col], errors="coerce")

    return df


def validate_dataframe(df):
    if df is None or df.empty:
        raise ValueError("Empty file uploaded.")

    df = normalize_dataframe(df)
    if "Name" not in df.columns:
        raise ValueError("Missing required column: Name.")

    subject_columns = [col for col in df.columns if col not in ["Name", "Total", "Average", "Grade", "Rank"]]
    if not subject_columns:
        raise ValueError("Missing required subject columns.")

    return df


def create_excel_file(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Processed Students")
        writer.save()
    output.seek(0)
    return output.read()


def build_analytics(df):
    subject_columns = [col for col in df.columns if col not in ["Rank", "Name", "Total", "Average", "Grade"]]
    grade_counts = df["Grade"].value_counts().reindex(["A", "B", "C", "Fail"], fill_value=0).to_dict()
    top_students = df.head(6)

    return {
        "totalStudents": int(len(df)),
        "topStudent": top_students.iloc[0]["Name"] if not top_students.empty else "N/A",
        "averageClassScore": round(float(df["Average"].mean()), 2) if not df.empty else 0,
        "gradeCounts": grade_counts,
        "gradeLabels": ["A", "B", "C", "Fail"],
        "gradeData": [int(grade_counts[label]) for label in ["A", "B", "C", "Fail"]],
        "topStudentLabels": top_students["Name"].astype(str).tolist(),
        "topStudentValues": top_students["Average"].round(2).tolist(),
        "subjectLabels": subject_columns,
        "subjectAverages": [round(float(df[col].mean()), 2) for col in subject_columns],
    }


@app.route("/")
def home():
    return send_from_directory(PUBLIC_DIR, "index.html")


@app.route("/process", methods=["POST"])
def process():
    if "file" not in request.files:
        return jsonify({"error": "No file selected."}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected."}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Wrong file type. Upload CSV, XLSX, DOCX, or PDF."}), 400

    try:
        df = load_dataframe(file)
        df = validate_dataframe(df)
        df = process_data(df)

        excel_bytes = create_excel_file(df)
        encoded = base64.b64encode(excel_bytes).decode("utf-8")

        response = {
            "filename": "processed_students.xlsx",
            "fileBytes": encoded,
            "analytics": build_analytics(df),
        }

        logger.info("File '%s' processed successfully with %d rows.", file.filename, len(df))
        return jsonify(response)

    except ValueError as error:
        logger.warning("Validation error processing file '%s': %s", file.filename, str(error))
        return jsonify({"error": str(error)}), 400
    except Exception as error:
        logger.exception("Server error processing file '%s'.", file.filename)
        return jsonify({"error": "Server processing error. Please check your file and try again."}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=os.environ.get("FLASK_DEBUG", "false").lower() == "true")